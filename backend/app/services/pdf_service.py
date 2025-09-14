"""
PDF Generation Service for Pactoria MVP
Generates PDF documents from contract content with proper formatting
"""

import io
from typing import Optional, Dict, Any
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.colors import black, grey
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE


class PDFService:
    """Service for generating PDF documents from contract content"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom paragraph styles for contract formatting"""
        
        # Contract title style
        self.styles.add(ParagraphStyle(
            name='ContractTitle',
            parent=self.styles['Title'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=black
        ))
        
        # Contract section heading
        self.styles.add(ParagraphStyle(
            name='SectionHeading',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=12,
            textColor=black
        ))
        
        # Contract body text
        self.styles.add(ParagraphStyle(
            name='ContractBody',
            parent=self.styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=12,
            alignment=TA_JUSTIFY,
            textColor=black
        ))
        
        # Contract clause
        self.styles.add(ParagraphStyle(
            name='ContractClause',
            parent=self.styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=8,
            leftIndent=20,
            alignment=TA_JUSTIFY,
            textColor=black
        ))
        
        # Footer style
        self.styles.add(ParagraphStyle(
            name='Footer',
            parent=self.styles['Normal'],
            fontSize=9,
            alignment=TA_CENTER,
            textColor=grey
        ))
    
    def generate_contract_pdf(
        self,
        title: str,
        content: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Generate a PDF document from contract content
        
        Args:
            title: Contract title
            content: Contract content (can include basic HTML-like formatting)
            metadata: Additional metadata like client info, dates, etc.
            
        Returns:
            PDF document as bytes
        """
        buffer = io.BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm,
            title=title
        )
        
        # Build document content
        story = []
        
        # Add title
        story.append(Paragraph(title, self.styles['ContractTitle']))
        story.append(Spacer(1, 20))
        
        # Add metadata table if provided
        if metadata:
            story.extend(self._build_metadata_table(metadata))
            story.append(Spacer(1, 20))
        
        # Add contract content
        story.extend(self._parse_content(content))
        
        # Add footer
        story.append(Spacer(1, 30))
        footer_text = f"Document generated on {datetime.now().strftime('%d %B %Y')} | Pactoria Contract Management"
        story.append(Paragraph(footer_text, self.styles['Footer']))
        
        # Build PDF
        doc.build(story)
        
        # Get PDF bytes
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        return pdf_bytes
    
    def _build_metadata_table(self, metadata: Dict[str, Any]) -> list:
        """Build a table with contract metadata"""
        data = []
        
        # Contract information
        if metadata.get('contract_type'):
            data.append(['Contract Type:', metadata['contract_type']])
        
        if metadata.get('client_name'):
            data.append(['Client:', metadata['client_name']])
            
        if metadata.get('supplier_name'):
            data.append(['Supplier:', metadata['supplier_name']])
            
        if metadata.get('contract_value') and metadata.get('currency'):
            value_str = f"{metadata['currency']} {metadata['contract_value']:,.2f}"
            data.append(['Contract Value:', value_str])
            
        if metadata.get('start_date'):
            data.append(['Start Date:', metadata['start_date']])
            
        if metadata.get('end_date'):
            data.append(['End Date:', metadata['end_date']])
        
        if not data:
            return []
        
        # Create table
        table = Table(data, colWidths=[3*cm, 10*cm])
        table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, grey),
            ('BACKGROUND', (0, 0), (0, -1), '#f0f0f0'),
        ]))
        
        return [table]
    
    def _parse_content(self, content: str) -> list:
        """Parse contract content and return formatted paragraphs"""
        story = []
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
                
            # Detect headings (lines that are all caps or start with numbers)
            if (para_text.isupper() and len(para_text) < 100) or \
               (para_text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) and '\n' not in para_text):
                story.append(Paragraph(para_text, self.styles['SectionHeading']))
            # Detect clauses (indented or numbered items)
            elif para_text.startswith(('a)', 'b)', 'c)', 'd)', 'e)', '(a)', '(b)', '(c)', 'i)', 'ii)', 'iii)')):
                story.append(Paragraph(para_text, self.styles['ContractClause']))
            else:
                # Regular paragraph
                story.append(Paragraph(para_text, self.styles['ContractBody']))
            
            story.append(Spacer(1, 6))
        
        return story
    
    def generate_contract_docx(
        self,
        title: str,
        content: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Generate a DOCX document from contract content
        
        Args:
            title: Contract title
            content: Contract content
            metadata: Additional metadata
            
        Returns:
            DOCX document as bytes
        """
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = 1  # Center alignment
        
        # Add metadata table if provided
        if metadata:
            doc.add_paragraph()  # Space
            self._add_docx_metadata_table(doc, metadata)
            doc.add_paragraph()  # Space
        
        # Add content
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
                
            # Detect headings
            if (para_text.isupper() and len(para_text) < 100) or \
               (para_text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) and '\n' not in para_text):
                doc.add_heading(para_text, level=2)
            else:
                para = doc.add_paragraph(para_text)
                para.alignment = 3  # Justify
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        
        return docx_bytes
    
    def _add_docx_metadata_table(self, doc: Document, metadata: Dict[str, Any]):
        """Add metadata table to DOCX document"""
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        metadata_items = []
        if metadata.get('contract_type'):
            metadata_items.append(('Contract Type:', metadata['contract_type']))
        if metadata.get('client_name'):
            metadata_items.append(('Client:', metadata['client_name']))
        if metadata.get('supplier_name'):
            metadata_items.append(('Supplier:', metadata['supplier_name']))
        if metadata.get('contract_value') and metadata.get('currency'):
            value_str = f"{metadata['currency']} {metadata['contract_value']:,.2f}"
            metadata_items.append(('Contract Value:', value_str))
        if metadata.get('start_date'):
            metadata_items.append(('Start Date:', metadata['start_date']))
        if metadata.get('end_date'):
            metadata_items.append(('End Date:', metadata['end_date']))
        
        for label, value in metadata_items:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(value)
            # Make first column bold
            row_cells[0].paragraphs[0].runs[0].bold = True


# Global service instance
pdf_service = PDFService()